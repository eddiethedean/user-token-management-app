from pathlib import Path
import math

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SHOT_DIR = ROOT / "page-screenshots"
CHUNK_DIR = ROOT / "screenshot-chunks"
OUTPUT = ROOT / "Data_Mover_Demo_Screenshots_v1.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(85, 96, 110)
LIGHT = "E8EEF5"
RULE = "B8C7D9"

PAGES = [
    ("1. Sign in", "01-sign-in.png",
     "Lead with trust and control: users can move data while their credentials remain encrypted, observable, and under their ownership."),
    ("2. Request access", "02-request-access.png",
     "Show a governed onboarding path: access begins with a verified government email and requires administrator review."),
    ("3. Password recovery", "03-password-recovery.png",
     "Reassure users that account recovery is simple and secure, using an eligible government email and a time-limited reset link."),
    ("4. Pipeline builder", "04-pipeline-builder.png",
     "Position Data Mover as an end-to-end workspace: configure a route, understand the transformation, and observe a realistic transfer from one screen."),
    ("5. Connections", "05-connections.png",
     "Demonstrate secure bring-your-own-connection management: each provider is configured once, encrypted at rest, and visibly health-checked."),
    ("6. Account settings", "06-account-settings.png",
     "Show that users retain control over their identity, password, sessions, and recent account activity in one clear account hub."),
    ("7. Users and invitations", "07-user-administration.png",
     "Convey deliberate access governance: administrators can provision, review, filter, invite, and disable application-managed identities."),
    ("8. Audit activity", "08-audit-history.png",
     "Establish accountability: security-relevant actions are recorded with time, outcome, source, and detail for operational review."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=180, bottom=120, end=180):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size, color=INK, bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])
    set_run(run, 9, MUTED)


def split_image(source, stem):
    CHUNK_DIR.mkdir(parents=True, exist_ok=True)
    with Image.open(source) as image:
        image = image.convert("RGB")
        max_height = 950
        if image.height <= max_height:
            out = CHUNK_DIR / f"{stem}-01.jpg"
            image.save(out, "JPEG", quality=94, optimize=True)
            return [out]
        chunks = []
        overlap = 28
        chunk_count = math.ceil((image.height - overlap) / (max_height - overlap))
        base_height = math.ceil((image.height + overlap * (chunk_count - 1)) / chunk_count)
        top = 0
        index = 1
        while top < image.height:
            bottom = min(top + base_height, image.height)
            crop = image.crop((0, top, image.width, bottom))
            out = CHUNK_DIR / f"{stem}-{index:02d}.jpg"
            crop.save(out, "JPEG", quality=94, optimize=True)
            chunks.append(out)
            if bottom == image.height:
                break
            top = bottom - overlap
            index += 1
        return chunks


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.top_margin = Inches(0.58)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.28)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for name, size, before, after in (("Title", 28, 0, 8), ("Heading 1", 18, 10, 6), ("Heading 2", 13, 8, 4)):
    style = styles[name]
    style.font.name = "Aptos Display" if name != "Heading 2" else "Aptos"
    style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = INK if name == "Title" else BLUE
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.text = "DATA MOVER  /  DEMO SCREENSHOT GUIDE"
set_run(header.runs[0], 8.5, MUTED, bold=True)
header.paragraph_format.space_after = Pt(0)

footer = section.footer.paragraphs[0]
add_page_number(footer)

# Editorial cover: the landscape orientation is a named screenshot-review override
# to the compact_reference_guide preset, preserving UI legibility.
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(52)
run = cover.add_run("DATA MOVER")
set_run(run, 11, BLUE, bold=True)
cover.paragraph_format.space_after = Pt(10)

title = doc.add_paragraph(style="Title")
title.add_run("Demo Screenshots")
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(18)
set_run(subtitle.add_run("Page-by-page product story and intended message"), 16, MUTED)

rule = doc.add_paragraph()
rule.paragraph_format.space_after = Pt(18)
p_pr = rule._p.get_or_add_pPr()
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "2E74B5")
p_bdr.append(bottom)
p_pr.append(p_bdr)

intro = doc.add_paragraph()
set_run(intro.add_run("Purpose. "), 11, INK, bold=True)
set_run(intro.add_run(
    "This guide captures the principal Data Mover screens and states the single idea each screen should communicate. "
    "Long pages continue across panels so interface text remains readable."
), 11, INK)

meta = doc.add_paragraph()
meta.paragraph_format.space_before = Pt(18)
set_run(meta.add_run("Version 1  ·  Demo environment  ·  August 18, 2026"), 10, MUTED)

for page_index, (heading, filename, message) in enumerate(PAGES):
    # Connections uses one complete desktop viewport rather than enlarged
    # scroll-page slices, preserving the full-width screen composition.
    if filename == "05-connections.png":
        viewport_path = CHUNK_DIR / "05-connections-full-screen.jpg"
        CHUNK_DIR.mkdir(parents=True, exist_ok=True)
        with Image.open(SHOT_DIR / filename) as image:
            image.convert("RGB").crop((0, 0, image.width, min(880, image.height))).save(
                viewport_path, "JPEG", quality=94, optimize=True
            )
        chunks = [viewport_path]
    else:
        chunks = split_image(SHOT_DIR / filename, Path(filename).stem)
    for chunk_index, chunk in enumerate(chunks):
        doc.add_page_break()
        if chunk_index == 0:
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.keep_with_next = True
            p.add_run(heading)

            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(9.45)
            cell = table.cell(0, 0)
            cell.width = Inches(9.45)
            set_cell_shading(cell, LIGHT)
            set_cell_margins(cell)
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            set_run(cp.add_run(message), 10.5, INK)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.space_before = Pt(2)
        else:
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.keep_with_next = True
            p.add_run(f"{heading} — continued")

        with Image.open(chunk) as im:
            ratio = im.height / im.width
        if filename == "05-connections.png":
            available_height = 4.7
        else:
            available_height = 5.58 if chunk_index == 0 else 6.7
        width = min(9.45, available_height / ratio)
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_before = Pt(0)
        pic.paragraph_format.space_after = Pt(0)
        picture = pic.add_run().add_picture(str(chunk), width=Inches(width))
        alt = heading if chunk_index == 0 else f"{heading}, continued panel {chunk_index + 1}"
        picture._inline.docPr.set("descr", f"Data Mover screenshot: {alt}")
        picture._inline.docPr.set("title", alt)

doc.core_properties.title = "Data Mover Demo Screenshots v1"
doc.core_properties.subject = "Page-by-page Data Mover product screenshots and intended messages"
doc.core_properties.author = "Data Mover"
doc.core_properties.keywords = "Data Mover, demo, screenshots, product story"
doc.save(OUTPUT)
print(OUTPUT)
